import logging
import re

logger = logging.getLogger(__name__)


def parse_document(file_path: str, file_type: str) -> str:
    """
    Extract plain text from a document file.

    Args:
        file_path: Absolute path to the file.
        file_type: One of 'pdf', 'doc', 'docx', 'txt'.

    Returns:
        Cleaned plain text.
    """
    ext = file_type.lower().lstrip(".")

    if ext == "pdf":
        return _parse_pdf(file_path)
    elif ext in ("doc", "docx"):
        return _parse_docx(file_path)
    elif ext == "txt":
        return _parse_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


def _parse_pdf(file_path: str) -> str:
    from pdfminer.high_level import extract_text

    try:
        text = extract_text(file_path)
        return _clean(text or "")
    except Exception as exc:
        logger.warning("PDF extraction error: %s", exc)
        return ""


def _parse_docx(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)
    return _clean("\n".join(paragraphs))


def _parse_txt(file_path: str) -> str:
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            with open(file_path, "r", encoding=encoding) as f:
                return _clean(f.read())
        except UnicodeDecodeError:
            continue
    # Final fallback: ignore undecodable bytes
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return _clean(f.read())


def _clean(text: str) -> str:
    """Normalize whitespace and strip noise."""
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse runs of blank lines to max 2
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Strip trailing whitespace per line
    lines = [line.rstrip() for line in text.splitlines()]
    return "\n".join(lines).strip()
